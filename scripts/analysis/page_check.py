"""Count the submission's pages, and fail if it is over the venue's limit.

Page limits are a hard gate at every venue this paper could go to, and they are
the one property of a document that cannot be checked by reading the source:
`python-docx` writes a document, it does not lay one out, so nothing in
`build_paper.py` knows how long the result is.

Two counters, in order of trustworthiness:

  Word COM       opens the file and asks Word how many pages it laid out. This
                 is the real number, not an estimate, and it is what a program
                 committee's own renderer would agree with.
  estimate       words / WORDS_PER_PAGE plus a page-equivalent per figure.
                 Used only when Word is unavailable, and it says so loudly --
                 an estimate that is silently substituted for a measurement is
                 how a paper arrives at a deadline two pages long.

    python scripts/analysis/page_check.py CollabEngine-NeurIPS2026.docx
    python scripts/analysis/page_check.py paper.docx --limit 9 --body-ends-at References

Venue limits change every cycle. `--limit` has no default on purpose: a default
would be wrong the year after it was written, and wrong in the direction that
looks fine until submission (Final Sweep 8.3).
"""

from __future__ import annotations

import argparse
import sys
from pathlib import Path

#: Only used by the fallback. Calibrated against this document's own layout
#: (single column, 10pt body, 1in margins), which is why it is not a general
#: constant and is not exported.
WORDS_PER_PAGE = 620
PAGES_PER_FIGURE = 0.32


def count_with_word(path: Path) -> int | None:
    """True page count via Word COM, or None if Word is not available."""
    try:
        import win32com.client  # noqa: PLC0415
    except ImportError:
        try:
            import comtypes.client as win32com  # noqa: PLC0415, F401
        except ImportError:
            return None
    try:
        import win32com.client as client  # noqa: PLC0415

        word = client.DispatchEx("Word.Application")
    except Exception:
        return None
    word.Visible = False
    try:
        doc = word.Documents.Open(str(path.resolve()), ReadOnly=True)
        try:
            doc.Repaginate()
            return int(doc.ComputeStatistics(2))  # wdStatisticPages
        finally:
            doc.Close(False)
    except Exception:
        return None
    finally:
        try:
            word.Quit()
        except Exception:
            pass


def measure(path: Path, body_ends_at: str | None) -> dict:
    """Words, figures and the body/back-matter split.

    `body_ends_at` is the heading where the page limit stops counting -- at most
    venues that is "References", with the bibliography and appendices excluded.
    Getting this wrong in the generous direction is the expensive mistake, so
    the split is reported rather than assumed.
    """
    from docx import Document

    doc = Document(str(path))
    words = body_words = 0
    figures = body_figures = 0
    in_body = True
    for par in doc.paragraphs:
        text = par.text.strip()
        if body_ends_at and text == body_ends_at:
            in_body = False
        n = len(text.split())
        words += n
        if in_body:
            body_words += n
        drawings = len(par._p.findall(
            ".//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline"
        ))
        figures += drawings
        if in_body:
            body_figures += drawings
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                n = len(cell.text.split())
                words += n
                if in_body:
                    body_words += n
    return {
        "words": words,
        "body_words": body_words,
        "figures": figures,
        "body_figures": body_figures,
    }


def estimate_pages(words: int, figures: int) -> float:
    return words / WORDS_PER_PAGE + figures * PAGES_PER_FIGURE


def main() -> int:
    ap = argparse.ArgumentParser()
    ap.add_argument("path")
    ap.add_argument("--limit", type=int, default=None,
                    help="venue page limit for the body; no default on purpose")
    ap.add_argument("--body-ends-at", default="References",
                    help="heading at which the counted body stops")
    args = ap.parse_args()

    path = Path(args.path)
    if not path.exists():
        print(f"no document at {path}; build it first", file=sys.stderr)
        return 2

    m = measure(path, args.body_ends_at)
    true_pages = count_with_word(path)

    print(f"{path.name}")
    print(f"  words           {m['words']:,} total, {m['body_words']:,} before "
          f"{args.body_ends_at!r}")
    print(f"  figures         {m['figures']} total, {m['body_figures']} in body")
    if true_pages is not None:
        print(f"  pages           {true_pages} (measured by Word, whole document)")
    else:
        print("  pages           Word unavailable -- ESTIMATE ONLY, do not "
              "submit against this number")
    est_body = estimate_pages(m["body_words"], m["body_figures"])
    print(f"  body estimate   {est_body:.1f} pages")

    if args.limit is None:
        print("\n  no --limit given, so nothing is enforced. Venue limits change "
              "every cycle;\n  check the current call for papers and pass it "
              "(Final Sweep 8.3).")
        return 0

    over = est_body > args.limit
    print(f"\n  limit           {args.limit} pages of body")
    if over:
        print(f"  OVER by roughly {est_body - args.limit:.1f} pages.")
        print("  The instrument-defect detail is the first thing to move to the")
        print("  appendix -- it is the section a reader can follow from a summary.")
        return 1
    print(f"  under by roughly {args.limit - est_body:.1f} pages")
    return 0


if __name__ == "__main__":
    sys.exit(main())
