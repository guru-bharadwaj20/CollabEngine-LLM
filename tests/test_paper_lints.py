"""The two build-time assertions that guard the submission.

Both exist because the defect they catch is *silent*. A bibliography whose
numbers have drifted still renders; a table reporting only significance still
reads as a finding. Neither shows up in a proofread, and this project shipped
one of each -- two references with invented titles, and tables that report a
p-value with no effect size beside it.

The build raises rather than warns. These tests assert that it does, and that
it does not raise on the shapes it should accept.
"""

from __future__ import annotations

import importlib.util
import sys
from pathlib import Path

import pytest

BUILD = Path(__file__).resolve().parents[1] / "scripts" / "analysis" / "build_paper.py"
pytest.importorskip("docx")


def _module():
    """Load build_paper.py by path -- scripts/ is not an importable package."""
    spec = importlib.util.spec_from_file_location("build_paper", BUILD)
    mod = importlib.util.module_from_spec(spec)
    sys.modules.setdefault("build_paper", mod)
    spec.loader.exec_module(mod)
    return mod


@pytest.fixture(scope="module")
def bp():
    return _module()


# ---------------------------------------------------------------------------
# no p-value without an effect size
# ---------------------------------------------------------------------------


@pytest.mark.parametrize("header", ["p", "*p*", "perm p", "TOST p", "Holm", "BH-FDR"])
def test_p_value_columns_are_recognised(bp, header):
    assert bp._is_p_column(header)


@pytest.mark.parametrize("header", ["tier", "solo", "team", "n", "model"])
def test_ordinary_columns_are_not_mistaken_for_p_values(bp, header):
    assert not bp._is_p_column(header)


@pytest.mark.parametrize(
    "header", ["*d*", "95% CI", "gap", "bound", "difference", "MDE", "sd"]
)
def test_effect_size_columns_are_recognised(bp, header):
    assert bp._is_effect_column(header)


def test_a_table_with_a_p_value_and_no_effect_size_is_refused(bp):
    from docx import Document

    with pytest.raises(AssertionError) as exc:
        bp.table(Document(), "a gate that reports only significance",
                 ["tier", "perm p"], [["medium", "0.269"]], number=1)
    assert "p-value" in str(exc.value)


def test_the_same_table_is_accepted_once_an_interval_is_added(bp):
    from docx import Document

    bp.table(Document(), "a gate that reports what it found",
             ["tier", "gap", "95% CI", "perm p"],
             [["medium", "+0.059", "[-0.03, +0.15]", "0.269"]], number=1)


def test_a_table_with_no_p_value_at_all_is_left_alone(bp):
    from docx import Document

    bp.table(Document(), "an inventory", ["arm", "n"], [["solo", "150"]], number=1)


# ---------------------------------------------------------------------------
# every citation resolves, every reference is cited
# ---------------------------------------------------------------------------


def _doc_with(bp, body: list[str]):
    from docx import Document

    doc = Document()
    for line in body:
        doc.add_paragraph(line)
    doc.add_paragraph("References")
    doc.add_paragraph("[1]  something that is only a reference entry")
    return doc


def test_a_citation_with_no_reference_is_refused(bp):
    doc = _doc_with(bp, ["As shown in [3], the gap is null."])
    with pytest.raises(AssertionError) as exc:
        bp.check_citations(doc, n_refs=2)
    assert "no reference" in str(exc.value)


def test_a_reference_that_is_never_cited_is_refused(bp):
    doc = _doc_with(bp, ["As shown in [1], the gap is null."])
    with pytest.raises(AssertionError) as exc:
        bp.check_citations(doc, n_refs=2)
    assert "never cited" in str(exc.value)


def test_grouped_citations_count_as_citing_each_member(bp):
    doc = _doc_with(bp, ["Several results agree [1, 2, 3]."])
    bp.check_citations(doc, n_refs=3)


def test_reference_entries_themselves_do_not_count_as_citations(bp):
    """The numbering in the reference list must not satisfy its own check.

    Without the stop at the References heading, a bibliography would cite
    itself and the check would pass on any document at all.
    """
    doc = _doc_with(bp, ["No citations in this body at all."])
    with pytest.raises(AssertionError):
        bp.check_citations(doc, n_refs=1)


# ---------------------------------------------------------------------------
# the paper is not the research log
# ---------------------------------------------------------------------------


def _prose(bp, *lines):
    from docx import Document

    doc = Document()
    for line in lines:
        doc.add_paragraph(line)
    return doc


def test_a_research_log_reference_in_the_body_is_refused(bp):
    doc = _prose(bp, "The cap artifact is documented in RESEARCH-LOG 4.14.")
    with pytest.raises(AssertionError) as exc:
        bp.check_register(doc)
    assert "research-log reference" in str(exc.value)


def test_first_person_singular_is_refused(bp):
    doc = _prose(bp, "I coded 40 messages by reading them.")
    with pytest.raises(AssertionError) as exc:
        bp.check_register(doc)
    assert "first person singular" in str(exc.value)


def test_a_calendar_date_in_prose_is_refused(bp):
    doc = _prose(bp, "Re-measured on 2026-08-13, the gap is negative.")
    with pytest.raises(AssertionError):
        bp.check_register(doc)


def test_a_diary_marker_is_refused(bp):
    doc = _prose(bp, "The run finished overnight and the arms disagree.")
    with pytest.raises(AssertionError):
        bp.check_register(doc)


def test_the_papers_own_section_numbers_are_not_flagged(bp):
    """A lint that cannot tell a self-reference from a log reference gets
    switched off within a week, so it does not try to."""
    bp.check_register(_prose(bp, "Read together with the asymmetry of section 4.2 (\u00a74.2)."))


def test_ordinary_academic_prose_passes(bp):
    bp.check_register(_prose(
        bp,
        "We report the artifact mechanism and a bounded negative result.",
        "Effects larger than 0.032 are excluded [1].",
        "Multi-agent LLM systems consume more generation than a single agent.",
    ))


def test_the_first_person_check_is_strict_and_that_is_deliberate(bp):
    """A standalone capital I fires even when it is not first person.

    Written down because it was found by this suite rather than in review: the
    sentence "the word AI contains a capital I" trips the lint. In academic
    prose a bare `I` is first person essentially always, so the rule stays
    strict and the false positive is accepted rather than weakened into a
    heuristic that lets the real thing through.
    """
    with pytest.raises(AssertionError):
        bp.check_register(_prose(bp, "The word AI contains a capital I here."))
