"""The double-blind build carries nothing that resolves to the authors.

`--anonymous` had existed for weeks without ever being run, which is the worst
state for a submission switch to be in: it looks like the anonymity question is
handled, and nothing has ever checked that it is. This suite runs it.

**What it covers.** The full document assembly -- title block, body, appendix
and the answered NeurIPS checklist appended by `checklist.py` -- is built with
`--anonymous` in memory, and every string a reviewer can read out of it is
grepped for the author name, the surname alone, the contact address, its local
part, the affiliation, the GitHub handle, any GitHub link, the repository name
and the lab and university names. The grep here is written out independently of
`build_paper.check_anonymous` rather than delegating to it, because a test that
only calls the lint proves the lint runs, not that the document is clean.

**What it cannot cover yet.** `numbers()` reads the corpus out of `runs/`, and
the corpus for the headline arms is not present in a checkout, so the build's
own numbers are stubbed here and the `.docx` is never written to disk. That
leaves two things unchecked: that a real corpus builds at all (an existing
concern, not an anonymity one), and the saved file's on-disk name. Neither can
leak an identifier that this suite would have caught -- the stubbed values are
floats, and the output filename is asserted separately from the argument
parser's own defaulting. When the corpus lands, the honest upgrade is to build
for real and grep the saved file.
"""

from __future__ import annotations

import importlib.util
import sys
from pathlib import Path

import pytest

ROOT = Path(__file__).resolve().parents[1]
BUILD = ROOT / "scripts" / "analysis" / "build_paper.py"
FIGURES = ROOT / "docs" / "figures" / "paper"
pytest.importorskip("docx")

AUTHOR = "Guru R Bharadwaj"
AFFILIATION = "Independent Researcher"
EMAIL = "gururb20@gmail.com"
HANDLE = "guru-bharadwaj20"

#: What the paper quotes from the corpus. Held here as a fixed dict rather than
#: a defaulting one so that a new quantity added to `numbers()` fails this file
#: loudly instead of silently rendering as zero.
STUB_NUMBERS = {
    "solo": 0.579, "team": 0.574, "three": 0.576,
    "solo_long": 0.516, "spread": 0.005, "n_total": 2900,
}


def _by_path(name: str, path: Path):
    """Load a `scripts/analysis` module by path -- scripts/ is not a package.

    Deliberately *not* done by putting that directory on `sys.path`, the way
    the builder does for itself when run as a script: it holds modules called
    `figures` and `rescore`, and importing the suite with those names ahead of
    the real package broke an unrelated test file two directories away. Under
    pytest the path stays as it is and the two modules are loaded by location.
    """
    spec = importlib.util.spec_from_file_location(name, path)
    mod = importlib.util.module_from_spec(spec)
    sys.modules.setdefault(name, mod)
    spec.loader.exec_module(mod)
    return mod


@pytest.fixture(scope="module")
def bp():
    return _by_path("build_paper", BUILD)


def _assemble(bp, anonymous: bool, acknowledgements: str = ""):
    """Everything `main` puts in the document, without the corpus or the save.

    Kept in step with `main` by hand, which is the cost of the corpus not being
    present. The order matters: the appendix and the checklist arrive after the
    lints the builder runs itself, and they are exactly where an identifier
    would appear without `build` ever seeing it.
    """
    if not FIGURES.is_dir():
        pytest.skip(f"paper figures absent: {FIGURES}")
    real, bp.numbers = bp.numbers, lambda: dict(STUB_NUMBERS)
    try:
        doc = bp.build(Path("unused.docx"), AUTHOR, AFFILIATION, EMAIL,
                       anonymous, FIGURES, acknowledgements)
    finally:
        bp.numbers = real
    bp.append_appendix(doc)
    checklist = _by_path("checklist", BUILD.parent / "checklist.py")
    checklist.append_checklist(doc, bp.para, bp.heading, bp.bullet, bp.runs)
    return doc


def _all_text(doc) -> str:
    """Body, tables, headers and footers, as one string to grep."""
    parts = [p.text for p in doc.paragraphs]
    for tbl in doc.tables:
        parts += [p.text for row in tbl.rows for c in row.cells
                  for p in c.paragraphs]
    for sec in doc.sections:
        for part in (sec.header, sec.footer,
                     sec.first_page_header, sec.first_page_footer):
            parts += [p.text for p in part.paragraphs]
    return "\n".join(parts)


@pytest.fixture(scope="module")
def anon_text(bp):
    return _all_text(_assemble(bp, anonymous=True))


# ---------------------------------------------------------------------------
# the lint proper: nothing identifying survives --anonymous
# ---------------------------------------------------------------------------


@pytest.mark.parametrize(
    "identifier",
    [AUTHOR, "Bharadwaj", EMAIL, "gururb20", AFFILIATION, HANDLE,
     "github.com", "https://github.com/guru-bharadwaj20/CollabEngine-LLM",
     "CollabEngine", "CCBD", "CDSAML", "PES University"],
)
def test_no_identifier_survives_the_anonymous_build(anon_text, identifier):
    assert identifier.lower() not in anon_text.lower(), (
        f"{identifier!r} is in the anonymous build"
    )


def test_the_builders_own_anonymity_check_passes_on_it(bp):
    bp.check_anonymous(_assemble(bp, anonymous=True), AUTHOR, AFFILIATION, EMAIL)


def test_the_anonymous_build_still_passes_the_citation_lint(bp):
    """The double-blind prose is prose, and obeys the same rules as the rest.

    §13's availability statement is different text in the two forms, so it is
    the one paragraph that could pass the builder's own lints in the preprint
    and fail them under `--anonymous` without anyone noticing.

    `check_register` is *not* asserted here, and the reason is a defect it has
    rather than a licence taken with the prose: it walks past the References
    heading instead of stopping at it, so an author initial in a reference entry
    ("and I. Mordatch") trips its first-person rule. That fires on both forms of
    the build equally and has nothing to do with anonymity, so it is left to be
    fixed where it belongs rather than worked around here.
    """
    doc = _assemble(bp, anonymous=True)
    bp.check_citations(doc, doc.n_references)


def test_the_docx_author_property_is_anonymous_too(bp):
    """A file property is not prose, and no proofread ever catches it."""
    doc = _assemble(bp, anonymous=True)
    assert AUTHOR not in (doc.core_properties.author or "")
    assert AUTHOR not in (doc.core_properties.last_modified_by or "")


def test_acknowledgements_are_dropped_however_they_are_passed(bp):
    """Acknowledgements name people and institutions by definition."""
    thanks = "The authors thank the CCBD-CDSAML lab for the compute."
    assert "CCBD" not in _all_text(_assemble(bp, True, acknowledgements=thanks))


def test_the_repo_link_is_replaced_rather_than_simply_deleted(anon_text):
    """Double-blind removes the link; it does not remove the commitment.

    Deleting the availability statement would leave the paper silent about
    whether the code exists at all, which reads worse to a reviewer than the
    conventional substitute. The anonymous form promises release on acceptance
    and says plainly that no anonymous mirror is being claimed.
    """
    assert "released publicly upon acceptance" in anon_text


# ---------------------------------------------------------------------------
# what --anonymous deliberately keeps, and what the preprint keeps
# ---------------------------------------------------------------------------


def test_the_gpu_model_is_deliberately_not_stripped(anon_text):
    """A judgement call, recorded here so it is revisited rather than assumed.

    "one 24 GB RTX 4500 Ada" names a card sold to anyone, not a machine in a
    named lab, and the conference asks for exactly this compute detail. Blurring
    it would cost a reviewer a reproducibility fact and buy no anonymity. If a
    hardware description ever names a lab or a cluster, this test is the wrong
    one and the description is the thing to change.
    """
    assert "RTX 4500 Ada" in anon_text


def test_the_conferences_own_links_are_not_treated_as_identifying(bp):
    """The checklist quotes `neurips.cc` URLs verbatim and must keep them."""
    from docx import Document

    doc = Document()
    doc.add_paragraph("See https://neurips.cc/public/EthicsGuidelines for more.")
    bp.check_anonymous(doc, AUTHOR, AFFILIATION, EMAIL)


def test_the_preprint_form_still_carries_the_byline_and_the_link(bp):
    """`--anonymous` is a switch, not a deletion.

    The failure mode this guards against is anonymising by removing the
    identifying material from the source, which passes every test above and
    leaves the preprint with no author and no code link.
    """
    text = _all_text(_assemble(bp, anonymous=False,
                               acknowledgements="The authors thank the lab."))
    assert AUTHOR in text
    assert EMAIL in text
    assert AFFILIATION in text
    assert bp.REPO_URL in text
    assert "The authors thank the lab." in text


# ---------------------------------------------------------------------------
# the check fails loudly, on every surface a reviewer can read
# ---------------------------------------------------------------------------


def test_an_identifier_in_the_body_is_refused(bp):
    from docx import Document

    doc = Document()
    doc.add_paragraph(f"Correspondence to {EMAIL}.")
    with pytest.raises(AssertionError) as exc:
        bp.check_anonymous(doc, AUTHOR, AFFILIATION, EMAIL)
    assert "contact address" in str(exc.value)


def test_an_identifier_in_a_table_cell_is_refused(bp):
    """Tables are where the compute description lives, and `doc.paragraphs`
    does not reach them -- the reason this check does not reuse the register
    check's walk."""
    from docx import Document

    doc = Document()
    doc.add_table(rows=1, cols=1).rows[0].cells[0].text = (
        "Serving: one 24 GB card in the CCBD lab"
    )
    with pytest.raises(AssertionError) as exc:
        bp.check_anonymous(doc, AUTHOR, AFFILIATION, EMAIL)
    assert "the lab" in str(exc.value)


def test_a_repository_link_in_a_footer_is_refused(bp):
    from docx import Document

    doc = Document()
    doc.sections[0].footer.paragraphs[0].text = f"Code: {bp.REPO_URL}"
    with pytest.raises(AssertionError) as exc:
        bp.check_anonymous(doc, AUTHOR, AFFILIATION, EMAIL)
    assert "footer" in str(exc.value)


def test_every_leak_is_reported_not_just_the_first(bp):
    """Fixing leaks one build at a time is how the second one ships."""
    from docx import Document

    doc = Document()
    doc.add_paragraph(AUTHOR)
    doc.add_paragraph(AFFILIATION)
    doc.add_paragraph(f"Code at {bp.REPO_URL}")
    with pytest.raises(AssertionError) as exc:
        bp.check_anonymous(doc, AUTHOR, AFFILIATION, EMAIL)
    message = str(exc.value)
    assert "the author name" in message
    assert "the affiliation" in message
    assert "the GitHub handle" in message


def test_a_short_acronym_does_not_fire_inside_an_ordinary_word(bp):
    """The lab acronyms are word-bounded on purpose.

    A lint that flags `CCBD` inside `ACCBDX` is a lint that gets switched off
    within a week, which is the same reasoning the register check records for
    its own section numbers.
    """
    from docx import Document

    doc = Document()
    doc.add_paragraph("The ACCBDX encoding and the PESUM statistic are unrelated.")
    bp.check_anonymous(doc, AUTHOR, AFFILIATION, EMAIL)
